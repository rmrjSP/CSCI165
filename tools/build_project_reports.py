from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(r"C:\CSCI165")


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def style_doc(doc: Document, title: str) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("CSCI 165 Project Report").italic = True


def add_section(doc: Document, heading: str, paragraphs: list[str]) -> None:
    doc.add_heading(heading, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_8queens() -> Path:
    doc = Document()
    style_doc(doc, "Incremental Evolutionary Algorithm for 8-Queens")
    summary = read_csv(ROOT / "8Queens" / "results" / "summary.csv")

    add_section(doc, "Motivation", [
        "The 8-queens problem is a classic example of combinatorial search and constraint satisfaction. It is simple to state, easy to visualize, and useful for demonstrating how search algorithms behave when the solution space contains many invalid or low-quality states.",
        "This domain is important because it provides a clean setting for comparing optimization strategies. In particular, it helps show how representation, selection pressure, and mutation design can strongly influence solution quality and runtime."
    ])

    add_section(doc, "Problem Statement", [
        "The goal of the project was to place eight queens on a standard chessboard so that no two queens attack each other. A valid solution must avoid row and diagonal conflicts.",
        "Beyond finding a valid board, the project also compared the performance of an evolutionary algorithm under different population sizes, mutation rates, and tournament sizes. The main question was which settings solve the problem most reliably and efficiently."
    ])

    add_section(doc, "Related Work and Background Material", [
        "The 8-queens problem has long been used as a benchmark for backtracking, local search, and evolutionary computation. Traditional exact approaches solve it with systematic search, while metaheuristics focus on finding strong solutions quickly without exploring every possible board.",
        "A key background idea for this project is representation. If queens are placed without restrictions, the search space becomes unnecessarily large because many boards repeat rows and are obviously poor candidates. By using a permutation encoding, each column contains exactly one queen and each row is used once, so the algorithm searches only row-distinct boards.",
        "The fitness function used in this project measures the number of non-attacking queen pairs. Since there are 28 total queen pairs on an 8-queen board, a perfect solution has fitness 28."
    ])

    add_section(doc, "Approach", [
        "The project used an evolutionary algorithm built around an incremental permutation encoding. Each chromosome is a permutation of the row indices 0 through 7, where the position in the permutation indicates the column and the value indicates the row.",
        "The algorithm begins with a randomly initialized population. Tournament selection chooses parents, cut-and-crossfill crossover preserves the permutation property, and swap mutation exchanges two row positions with a fixed probability. Elitism carries the best individuals forward unchanged, which helps preserve high-quality boards once they are found.",
        "This approach was chosen because it reduces wasted search, keeps candidate solutions valid with respect to row uniqueness, and provides a natural framework for testing how parameter changes affect convergence."
    ])

    add_section(doc, "Experiment Setup", [
        "The evaluation used 30 runs per configuration. All runs used elitism equal to 2, cut-and-crossfill crossover, and a maximum of 1000 generations.",
        "Three parameter sweeps were tested:"
    ])
    add_bullets(doc, [
        "Population size: 20, 50, and 100 with mutation rate 0.10 and tournament size 3.",
        "Mutation rate: 0.05, 0.10, and 0.20 with population size 100 and tournament size 3.",
        "Tournament size: 2, 3, and 5 with population size 100 and mutation rate 0.10."
    ])
    doc.add_paragraph("The main metrics were success rate, mean best fitness, average generation when a perfect solution was found, and runtime in milliseconds.")

    add_section(doc, "Results and Discussion", [
        "The results showed that the evolutionary algorithm performed very well on this problem when the population size was moderate or large. Every configuration except one solved the problem in all 30 runs.",
        f"The weakest configuration was population size 20, which achieved a success rate of {summary[0]['success_rate']} with mean fitness {summary[0]['mean_fitness']}. In contrast, population sizes 50 and 100 both reached {summary[1]['success_rate']} and {summary[2]['success_rate']} success respectively, with perfect mean fitness of 28.0.",
        "Among the mutation settings, all tested values solved the problem in every run, so mutation was not a major source of instability within the tested range. Tournament size also produced perfect success in all tested cases, although tournament size 5 required more generations on average than sizes 2 or 3, which suggests stronger selection pressure may reduce diversity too quickly.",
        "Overall, the project demonstrates that the representation was the most important design decision. Once the search space was restricted to row-distinct boards, the evolutionary algorithm was able to find perfect solutions consistently under many parameter settings."
    ])

    add_section(doc, "Conclusion and Future Work Directions", [
        "The main contribution of this project is a simple but effective evolutionary formulation of the 8-queens problem using permutation encoding. The method reliably found perfect solutions and showed that population size had the clearest effect on success and speed.",
        "The best practical settings in this experiment used larger populations with moderate mutation and tournament sizes. These settings solved the problem quickly and consistently across repeated trials.",
        "Future work could test larger N-queens instances, compare this approach with hill climbing or simulated annealing, and explore other crossover or mutation operators to understand how well the method scales beyond the 8-queen case."
    ])

    out = ROOT / "8Queens" / "8Queens_report.docx"
    doc.save(out)
    return out


def build_tsp() -> Path:
    doc = Document()
    style_doc(doc, "TSP with Hill Climbing, Simulated Annealing, and Threshold Accepting")
    summary = {row["algorithm"]: row for row in read_csv(ROOT / "TSP" / "results" / "summary.csv")}

    add_section(doc, "Motivation", [
        "The Traveling Salesman Problem is one of the most well-known optimization problems in computer science. It models the task of visiting a set of locations while minimizing travel cost, which connects directly to routing, logistics, scheduling, manufacturing, and circuit design.",
        "The domain is important because exact optimization becomes difficult as the number of cities grows. For that reason, practical work often depends on heuristic and metaheuristic methods that can find strong solutions without exhaustive search."
    ])

    add_section(doc, "Problem Statement", [
        "This project addressed a fixed 100-city Traveling Salesman Problem instance. The task was to find a short closed tour that visits every city exactly once and then returns to the starting city.",
        "The main objective was not only to compute a good tour, but also to compare three local-search strategies: hill climbing, simulated annealing, and threshold accepting. The comparison focused on which method finds the shortest tours and how strongly acceptance strategy affects performance."
    ])

    add_section(doc, "Related Work and Background Material", [
        "TSP is a standard benchmark in optimization and artificial intelligence because it is easy to describe but computationally challenging to solve exactly. This has made it a common testbed for local search, stochastic search, and hybrid metaheuristics.",
        "A tour in TSP can be represented as a permutation of city indices. The quality of a solution is the total Euclidean distance of the closed path. Local search methods improve solutions by repeatedly applying neighborhood operations that slightly modify the current route.",
        "A standard neighborhood operator for TSP is 2-opt, which reverses a segment of the route. This move is often much stronger than a simple swap because it can remove route crossings and lead to large cost reductions."
    ])

    add_section(doc, "Approach", [
        "The project implemented three related search strategies on the same TSP instance and with the same 2-opt neighborhood. Hill climbing accepted only strictly improving moves. A restart variant repeated hill climbing from multiple random initial routes to reduce the chance of getting trapped early.",
        "Simulated annealing introduced probabilistic acceptance of worse moves according to a temperature schedule. This allows the search to escape local minima, especially early in the run when the temperature is high. Two cooling schedules were tested: a faster one and a slower one.",
        "Threshold accepting used a deterministic alternative. Instead of accepting worse moves probabilistically, it accepted moves whose cost increase was within a gradually decreasing threshold. This method keeps some exploratory behavior while remaining simpler than simulated annealing."
    ])

    add_section(doc, "Experiment Setup", [
        "The experiment used a fixed dataset of 100 cities with random two-dimensional coordinates. All algorithms used the same evaluation budget of 100,000 neighbor evaluations and were repeated for 30 runs.",
        "The configurations tested were:"
    ])
    add_bullets(doc, [
        "HC_plain: hill climbing with one start.",
        "HC_restart: hill climbing with five random restarts.",
        "SA_fast: simulated annealing with T0 = 1000 and cooling alpha = 0.9990.",
        "SA_slow: simulated annealing with T0 = 1000 and cooling alpha = 0.9999.",
        "TA: threshold accepting with initial threshold 50.0 and 200 rounds."
    ])
    doc.add_paragraph("The main metrics were mean tour cost, standard deviation, best and worst final cost, and mean runtime.")

    add_section(doc, "Results and Discussion", [
        f"The strongest overall performer was SA_slow, which achieved the best mean cost of {summary['SA_slow']['mean_cost']} and the best observed route cost of {summary['SA_slow']['best_cost']}. This supports the idea that slower cooling preserves exploration long enough to escape local minima effectively.",
        f"Hill climbing was competitive, especially with restarts. HC_plain produced a mean cost of {summary['HC_plain']['mean_cost']}, while HC_restart improved this to {summary['HC_restart']['mean_cost']}. The restart mechanism helped, but it still did not outperform slow simulated annealing.",
        f"SA_fast performed the worst on average with mean cost {summary['SA_fast']['mean_cost']}. This suggests that cooling too quickly reduced the algorithm's ability to explore and caused it to behave more like greedy local search before it had located a strong basin.",
        f"Threshold accepting finished between hill climbing and the weaker annealing schedule, with mean cost {summary['TA']['mean_cost']}. Its results show that deterministic acceptance of slightly worse moves can improve exploration, but in this experiment it did not match the strongest stochastic schedule.",
        "Overall, the results highlight that acceptance policy and cooling behavior matter as much as the neighborhood operator. Since every method used 2-opt, most of the performance difference came from how each algorithm balanced greediness and exploration."
    ])

    add_section(doc, "Conclusion and Future Work Directions", [
        "This project showed that slow simulated annealing was the most effective method among the tested approaches for the fixed 100-city TSP instance. Hill climbing was fast and competitive, especially with restarts, but it remained more vulnerable to local minima.",
        "The project's main contribution is a controlled comparison in which all algorithms used the same city instance, neighborhood operator, and evaluation budget. That made it possible to isolate the effect of the acceptance mechanism itself.",
        "Future work could test larger city sets, tune threshold schedules and restart counts more carefully, compare swap against 2-opt more directly, or extend the study to additional metaheuristics such as genetic algorithms or ant-colony optimization."
    ])

    out = ROOT / "TSP" / "TSP_report.docx"
    doc.save(out)
    return out


def build_rastrigin() -> Path:
    doc = Document()
    style_doc(doc, "Gradient Descent vs Simulated Annealing on the Rastrigin Function")
    summary = {row["algorithm"]: row for row in read_csv(ROOT / "Rastrigin" / "results" / "summary.csv")}

    add_section(doc, "Motivation", [
        "Optimization is central to machine learning, engineering design, scientific computing, and many decision-making systems. A major challenge is that real objective functions often contain many local minima, plateaus, or noisy gradients.",
        "The Rastrigin function is important as a benchmark because it is multimodal and has a known global minimum. This makes it a useful test for understanding when local methods succeed, when they fail, and whether stochastic exploration helps."
    ])

    add_section(doc, "Problem Statement", [
        "The project studied minimization of the two-dimensional Rastrigin function over the domain [-5.12, 5.12] for both coordinates. The goal was to find or closely approach the global minimum at x = (0, 0), where f(x) = 0.",
        "The main comparison was between three gradient descent variants and simulated annealing. The question was which method would most reliably reach the global optimum and how the search behavior would differ across deterministic and stochastic strategies."
    ])

    add_section(doc, "Related Work and Background Material", [
        "Rastrigin is a standard benchmark in continuous optimization because its oscillatory cosine terms create many local minima around a broad quadratic structure. This makes it more difficult than a simple convex bowl.",
        "Gradient descent methods rely on local gradient information and often converge quickly when the landscape is smooth or when the starting point is favorable. However, in multimodal landscapes they can become trapped in local minima or behave poorly under unsuitable learning rates.",
        "Simulated annealing provides a different strategy by allowing occasional uphill moves. This stochastic behavior can help the search leave local traps, although performance depends heavily on the temperature schedule and perturbation size."
    ])

    add_section(doc, "Approach", [
        "The project implemented three gradient descent variants: fixed learning rate, decaying learning rate, and momentum-based descent. All methods evaluated the same analytical gradient of the Rastrigin function.",
        "The fixed-rate method used a constant step size. The decaying-rate method gradually reduced the learning rate over time, which was intended to improve stability and fine-grained convergence near a good solution. The momentum method accumulated past gradient information to smooth updates and potentially move more efficiently across the landscape.",
        "Simulated annealing used random perturbations within a bounded step radius and accepted worse moves with probability based on the current temperature. This provided a direct contrast between local deterministic descent and broader stochastic exploration."
    ])

    add_section(doc, "Experiment Setup", [
        "The experiment used 30 shared random starting points sampled uniformly from the valid domain. Every algorithm ran with a maximum of 5000 iterations so the comparison would be consistent.",
        "The tested settings were:"
    ])
    add_bullets(doc, [
        "GD_Fixed: alpha = 0.01.",
        "GD_Decaying: alpha0 = 0.1 with decay = 0.001.",
        "GD_Momentum: alpha = 0.01 and beta = 0.9.",
        "SA: T0 = 10.0, alpha = 0.995, T_min = 1e-4, and step radius = 0.5."
    ])
    doc.add_paragraph("A run counted as successful if it achieved f(x) < 1e-3. The main metrics were best objective value, success rate, distance to the optimum, and runtime.")

    add_section(doc, "Results and Discussion", [
        f"The clear winner was the decaying gradient descent variant. It achieved a success rate of {summary['GD_Decaying']['success_rate']} with mean objective value {summary['GD_Decaying']['mean_f']}, showing that gradual reduction of the learning rate matched this landscape very well.",
        f"Fixed gradient descent and momentum both struggled. GD_Fixed had mean objective value {summary['GD_Fixed']['mean_f']} with success rate {summary['GD_Fixed']['success_rate']}, while GD_Momentum had mean objective value {summary['GD_Momentum']['mean_f']} with success rate {summary['GD_Momentum']['success_rate']}. These results indicate that both methods were often trapped away from the global optimum.",
        f"Simulated annealing performed better than the weaker gradient methods in terms of exploration and best-case closeness to the optimum, with best value {summary['SA']['best_f']}, but it still had success rate {summary['SA']['success_rate']} under the strict threshold used in this study. This suggests that the chosen annealing schedule was not strong enough to consistently finish at the exact global basin within the given budget.",
        "The main lesson from the experiment is that exploration alone is not enough; the search schedule must be tuned to the landscape. In this case, the decaying learning rate provided the best balance between global progress early and stable convergence late."
    ])

    add_section(doc, "Conclusion and Future Work Directions", [
        "The main contribution of this project is a direct comparison between local gradient-based optimization and stochastic search on a multimodal benchmark. Among the tested methods, decaying gradient descent was the most reliable and achieved near-perfect performance on every run.",
        "The results also show that simulated annealing can approach the optimum closely, but parameter tuning matters greatly. The fixed learning rate and momentum variants were not robust enough for this benchmark under the chosen settings.",
        "Future work could test more aggressive annealing schedules, alternative step sizes and momentum values, higher-dimensional Rastrigin problems, or other continuous benchmarks to determine whether the same ranking holds more generally."
    ])

    out = ROOT / "Rastrigin" / "Rastrigin_report.docx"
    doc.save(out)
    return out


def main() -> None:
    outputs = [build_8queens(), build_tsp(), build_rastrigin()]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
